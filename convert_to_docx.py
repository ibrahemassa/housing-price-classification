#!/usr/bin/env python3
"""
Convert Markdown report to DOCX format.
Requires: pip install python-docx markdown
"""

try:
    import re

    from docx import Document
except ImportError:
    print("Installing required packages...")
    import subprocess

    subprocess.check_call(["pip", "install", "python-docx", "markdown"])
    import re

    from docx import Document


def markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX."""
    doc = Document()

    with open(md_file, encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Headers
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            p = doc.add_heading(line[5:], level=4)
        # Screenshot placeholders
        elif "[SCREENSHOT PLACEHOLDER:" in line:
            p = doc.add_paragraph()
            p.add_run("[SCREENSHOT PLACEHOLDER]").bold = True
            i += 1
            if i < len(lines) and lines[i].strip().startswith("*"):
                caption = lines[i].strip().lstrip("*").strip()
                p.add_run(f": {caption}").italic = True
        # Bold text (for labels)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip("*")).bold = True
        # Code blocks
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph("".join(code_lines))
                p.style = "Intense Quote"
        # Lists
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        # Tables
        elif "|" in line and not line.startswith("|---"):
            if i > 0 and "|" in lines[i - 1] and "---" not in lines[i - 1]:
                # Table row
                cells = [c.strip() for c in line.split("|") if c.strip()]
                if not hasattr(doc, "_current_table") or doc._current_table is None:
                    doc._current_table = doc.add_table(rows=1, cols=len(cells))
                    doc._current_table.style = "Light Grid Accent 1"
                    row = doc._current_table.rows[0]
                else:
                    row = doc._current_table.add_row()
                for j, cell_text in enumerate(cells):
                    row.cells[j].text = cell_text
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)

        i += 1

    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")


if __name__ == "__main__":
    import sys

    md_file = "PROJECT_REPORT_SIMPLE.md"
    docx_file = "PROJECT_REPORT.docx"

    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    if len(sys.argv) > 2:
        docx_file = sys.argv[2]

    markdown_to_docx(md_file, docx_file)

